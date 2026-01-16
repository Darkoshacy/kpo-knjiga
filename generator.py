from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd

def set_cell_border(cell, **kwargs):
    """
    Postavlja border za ćeliju tabele u Word-u.
    kwargs može biti: top, bottom, left, right
    Svaki sa vrednošću dict sa 'sz' (size), 'val' (style), 'color'
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    # Kreiranje bordure elementa
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            border_el = OxmlElement(f'w:{edge}')
            border_el.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            border_el.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            border_el.set(qn('w:space'), '0')
            border_el.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(border_el)
    
    tcPr.append(tcBorders)

def generate_kpo_docx(df, output_path):
    """
    Generiše KPO knjigu u DOCX formatu na osnovu DataFrame-a.
    
    Args:
        df: pandas DataFrame sa kolonama ['Datum', 'Klijent', 'Prihod']
        output_path: putanja gde će biti sačuvan DOCX fajl
    """
    doc = Document()
    
    # Podešavanje margina
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Ekstrakcija godine iz podataka
    godina = "2024"  # Default
    if not df.empty and 'Datum' in df.columns:
        try:
            # Pokušaj da izvučeš godinu iz prvog datuma
            prvi_datum = str(df.iloc[0]['Datum'])
            if '.' in prvi_datum:
                godina = prvi_datum.split('.')[-1].strip()
        except:
            pass
    
    # Zaglavlje sa podacima obveznika
    header_text = f"""Obrazac KPO

PIB 108404606
MBR: 63429058
Obveznik  Darko Petrović
Firma-radnja  OBUKA NA RAČUNARU I EDUKACIJA MEXICOM INFORMATIKA ĆUPRIJA
Sedište Šumadijska 20, 35230 Ćuprija
Šifra poreskog obveznika 108404606
Šifra delatnosti 8559

KNJIGA O OSTVARENOM PROMETU
PAUŠALNO OPOREZOVANIH OBVEZNIKA ZA {godina}. GODINU"""
    
    header_para = doc.add_paragraph(header_text)
    for run in header_para.runs:
        run.font.size = Pt(11)
    
    doc.add_paragraph()  # Razmak
    doc.add_paragraph()  # Dodatni razmak
    
    # Kreiranje tabele: Redni broj | Datum | Komitent | od prodaje | usluge | Svega (4+5)
    num_rows = len(df) + 2  # +2 za dva reda zaglavlja
    table = doc.add_table(rows=num_rows, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Podešavanje širine kolona
    table.columns[0].width = Inches(0.4)   # Redni broj - upola uže
    table.columns[1].width = Inches(0.9)   # Datum
    table.columns[2].width = Inches(3.2)   # Komitent - najšire
    table.columns[3].width = Inches(1.0)   # od prodaje
    table.columns[4].width = Inches(1.0)   # usluge
    table.columns[5].width = Inches(1.2)   # Svega
    
    # Prvi red zaglavlja
    header_row1 = table.rows[0].cells
    header_row1[0].text = 'Redni broj'
    header_row1[1].text = 'Datum'
    header_row1[2].text = 'Komitent'
    
    # Merge ćelija 3 i 4 za "Prihod od delatnosti"
    merged_prihod = header_row1[3].merge(header_row1[4])
    merged_prihod.text = 'Prihod od delatnosti'
    
    header_row1[5].text = 'Svega prihod od delatnosti\n(4 + 5)'
    
    # Drugi red zaglavlja
    header_row2 = table.rows[1].cells
    header_row2[0].text = '1'
    header_row2[1].text = '2'
    header_row2[2].text = '3'
    header_row2[3].text = 'od prodaje\n4'
    header_row2[4].text = 'usluge\n5'
    header_row2[5].text = '6'
    
    # Formatiranje oba reda zaglavlja
    for row_idx in [0, 1]:
        for cell in table.rows[row_idx].cells:
            # Formatiranje teksta
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
            
            # Pozadina za zaglavlje
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D9E1F2')  # Svetlo plava
            cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Popunjavanje redova podataka
    for table_idx, (df_idx, row) in enumerate(df.iterrows(), start=2):  # start=2 jer su prva dva reda zaglavlje
        row_cells = table.rows[table_idx].cells
        
        # Redni broj (indeks iz DataFrame-a)
        row_cells[0].text = str(df_idx)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Datum
        row_cells[1].text = str(row['Datum'])
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Komitent (naziv klijenta)
        row_cells[2].text = str(row['Klijent'])
        
        # od prodaje (prazno - svi prihodi su usluge)
        row_cells[3].text = ''
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # usluge (prihod)
        prihod_value = row['Prihod']
        row_cells[4].text = f"{prihod_value:,.2f}".replace(',', ' ')
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Svega (isti kao usluge)
        row_cells[5].text = f"{prihod_value:,.2f}".replace(',', ' ')
        row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Veličina fonta za redove podataka
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    # Dodavanje ukupnog zbira na kraju
    doc.add_paragraph()  # Razmak
    
    total_prihod = df['Prihod'].sum()
    total_paragraph = doc.add_paragraph()
    total_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_run = total_paragraph.add_run(f'UKUPAN PRIHOD: {total_prihod:,.2f} RSD'.replace(',', ' '))
    total_run.font.size = Pt(12)
    total_run.font.bold = True
    total_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Potpis i datum
    doc.add_paragraph()
    doc.add_paragraph()
    
    signature_line = doc.add_paragraph('_' * 40)
    signature_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    signature_text = doc.add_paragraph('(Potpis odgovornog lica)')
    signature_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature_text.runs[0].font.italic = True
    signature_text.runs[0].font.size = Pt(9)
    
    # Čuvanje dokumenta
    doc.save(output_path)
    print(f"✓ KPO knjiga uspešno kreirana: {output_path}")
